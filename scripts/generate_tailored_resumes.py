"""
Generate 4 tailored .docx resumes for targeted job applications.
Uses python-docx to create professional, 1-page resumes.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "06-Drafts", "outreach", "resumes"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Shared helpers ──────────────────────────────────────────────────────────

CONTACT_LINE = (
    "kdammilare33@gmail.com  |  Lagos, Nigeria (open to relocation)  |  "
    "github.com/Qui-vira"
)

SKILLS = {
    "data_analysis": "Python (pandas, NumPy, scikit-learn), SQL, PostgreSQL, Supabase",
    "stats": "Bayesian inference, probability calibration, backtesting, A/B testing, cohort analysis, time-series",
    "viz": "Tableau (learning), Power BI (learning), Superset, TypeScript/React dashboards",
    "eng": "FastAPI, Next.js, Docker, Railway, REST APIs, GitHub CI/CD",
    "market": "Binance API, Twelve Data API, Polymarket event contracts, on-chain data",
}

EDUCATION = [
    "Data Analysis Track (active) - SQL, Tableau, Power BI, Python/ML via Coursera",
    "Self-directed research in prediction markets, on-chain analytics, event contract calibration",
]


def make_doc():
    """Create a Document with tight margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.space_before = Pt(0)
    return doc


def add_name(doc, name="DAMILARE SAMUEL KEHINDE-DAVID"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_contact(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(CONTACT_LINE)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_section_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.underline = True


def add_experience_block(doc, title, org, date, bullets):
    # Title line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    if org:
        run2 = p.add_run(f"  |  {org}")
        run2.font.size = Pt(10)
        run2.font.name = "Calibri"
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if date:
        run3 = p.add_run(f"  |  {date}")
        run3.font.size = Pt(9)
        run3.font.name = "Calibri"
        run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    # Bullets
    for b in bullets:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.left_indent = Inches(0.25)
        run = bp.add_run(f"- {b}")
        run.font.size = Pt(9)
        run.font.name = "Calibri"


def add_bullet_list(doc, items, indent=0.25):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(f"- {item}")
        run.font.size = Pt(9)
        run.font.name = "Calibri"


def add_skills_line(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"


def add_profile(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"


# ── Experience block definitions ────────────────────────────────────────────

EXP_SIGNAL = {
    "title": "Crypto/Forex Signal Scoring Engine",
    "org": "Quivira Signal Bot",
    "date": "2025 - Present",
}

EXP_POLYMARKET = {
    "title": "Polymarket Prediction Bot",
    "org": "Independent Research",
    "date": "2025 - Present",
}

EXP_SUBSCRIPTION = {
    "title": "Subscription Analytics and Growth",
    "org": "Hustler's Krib Signal",
    "date": "2025 - Present",
}

EXP_PHARMA = {
    "title": "Product and Revenue Analytics",
    "org": "PharmaOS",
    "date": "2025 - Present",
}


# ── Resume 1: PrimeIntellect - Founding GTM Lead ───────────────────────────

def build_primeintellect():
    doc = make_doc()
    add_name(doc)
    add_subtitle(doc, "GTM Strategy  |  AI/ML Growth  |  Data-Driven Revenue Operations")
    add_contact(doc)

    add_section_header(doc, "Profile")
    add_profile(doc,
        "GTM strategist with hands-on experience building AI/ML data pipelines that serve as "
        "revenue-generating products. Built and scaled a paid signal subscription from zero to "
        "recurring revenue through data-driven funnel optimization, partnership outreach, and "
        "community-led growth. Translates complex technical concepts into clear, compelling "
        "messaging that drives sales enablement and user acquisition. Familiar with CRM tooling "
        "(HubSpot) and revenue operations workflows."
    )

    add_section_header(doc, "Experience")

    add_experience_block(doc,
        EXP_SIGNAL["title"], EXP_SIGNAL["org"], EXP_SIGNAL["date"],
        [
            "Built full AI/ML data pipeline: ingestion, indicator computation, Bayesian scoring, threshold filtering across Binance futures + forex pairs",
            "Designed product positioning and GTM strategy for signal bot as paid subscription service",
            "Stack: Python (pandas, NumPy), SQL, REST APIs, Docker, Railway - end-to-end pipeline ownership",
            "All signals logged to PostgreSQL with outcome tracking for data-driven performance marketing",
        ])

    add_experience_block(doc,
        EXP_SUBSCRIPTION["title"], EXP_SUBSCRIPTION["org"], EXP_SUBSCRIPTION["date"],
        [
            "Drove GTM execution: pricing experiments ($49/wk to $1,499 lifetime), funnel optimization, referral attribution (30% commission program)",
            "Tracked subscriber LTV, churn, cohort retention to inform revenue operations decisions",
            "Built signal performance dashboards in PostgreSQL to support sales enablement and partnership pitches",
        ])

    add_experience_block(doc,
        EXP_POLYMARKET["title"], EXP_POLYMARKET["org"], EXP_POLYMARKET["date"],
        [
            "Designed 7-factor Bayesian scoring system for calibrated probability estimates - demonstrates AI/ML research depth",
            "Factor analysis and weight auditing against historical resolution data",
            "Outputs structured predictions with confidence intervals, logged to Supabase PostgreSQL",
        ])

    add_experience_block(doc,
        EXP_PHARMA["title"], EXP_PHARMA["org"], EXP_PHARMA["date"],
        [
            "Built Super Admin dashboard: revenue per pharmacy, subscription lifecycle, tenant-level usage analytics",
            "Designed data model: 8,500+ drug products, patient routing, prescription history, billing",
        ])

    add_section_header(doc, "Domain Fit")
    add_bullet_list(doc, [
        "Deep AI infrastructure understanding from building ML-powered data pipelines that serve as the core product",
        "Grew paid signal community from 0 to recurring subscribers through partnerships, content, and data-driven GTM",
        "Experience translating technical complexity (Bayesian models, on-chain data) into clear value propositions for non-technical buyers",
    ])

    add_section_header(doc, "Technical Skills")
    add_skills_line(doc, "Data & AI/ML", SKILLS["data_analysis"])
    add_skills_line(doc, "Statistics", SKILLS["stats"])
    add_skills_line(doc, "Engineering", SKILLS["eng"])
    add_skills_line(doc, "GTM Tools", "HubSpot (familiar), Notion, PostgreSQL dashboards, Typefully, automated distribution pipelines")

    add_section_header(doc, "Education")
    add_bullet_list(doc, EDUCATION)

    path = os.path.join(OUTPUT_DIR, "Damilare_Kehinde-David_Resume_PrimeIntellect_GTM.docx")
    doc.save(path)
    print(f"  Saved: {path}")


# ── Resume 2: Kalshi - Growth Marketing Manager ────────────────────────────

def build_kalshi():
    doc = make_doc()
    add_name(doc)
    add_subtitle(doc, "Growth Marketing  |  Prediction Markets  |  Data-Driven Acquisition")
    add_contact(doc)

    add_section_header(doc, "Profile")
    add_profile(doc,
        "Performance-minded growth marketer with direct prediction market experience and deep SQL "
        "proficiency. Built a Polymarket prediction bot from scratch, ran pricing experiments across "
        "a paid trading community, and optimized conversion funnels from free to premium tiers. "
        "Combines hands-on data analysis with experimentation rigor to drive user acquisition, "
        "attribution modeling, and funnel optimization in fintech environments."
    )

    add_section_header(doc, "Experience")

    add_experience_block(doc,
        EXP_POLYMARKET["title"], EXP_POLYMARKET["org"], EXP_POLYMARKET["date"],
        [
            "Built Polymarket prediction bot: 7-factor Bayesian scoring system for calibrated probability estimates across political, economic, cultural categories",
            "Factor analysis and weight auditing against historical resolution data for continuous model improvement",
            "Active prediction market participant with deep understanding of event contract mechanics and regulatory landscape",
            "Outputs structured predictions with confidence intervals, logged to Supabase PostgreSQL for performance tracking",
        ])

    add_experience_block(doc,
        EXP_SUBSCRIPTION["title"], EXP_SUBSCRIPTION["org"], EXP_SUBSCRIPTION["date"],
        [
            "Ran A/B testing on pricing tiers ($49/wk to $1,499 lifetime) to optimize conversion rate and subscriber LTV",
            "Built attribution model tracking referral sources (30% commission program), paid acquisition, organic funnel performance",
            "Cohort retention analysis and churn prediction to inform user acquisition strategy",
            "Signal performance dashboards in PostgreSQL driving experimentation roadmap",
        ])

    add_experience_block(doc,
        EXP_SIGNAL["title"], EXP_SIGNAL["org"], EXP_SIGNAL["date"],
        [
            "Full data pipeline: ingestion, indicator computation, Bayesian scoring, threshold filtering across multiple markets",
            "Statistical analysis on win rates and outcome tracking - SQL-heavy analytical workflow",
            "Stack: Python (pandas, NumPy), SQL, REST APIs, Docker, Railway",
        ])

    add_experience_block(doc,
        EXP_PHARMA["title"], EXP_PHARMA["org"], EXP_PHARMA["date"],
        [
            "Revenue analytics dashboard: subscription lifecycle, tenant-level usage, conversion tracking",
            "Data model design: 8,500+ products, patient routing, billing analytics",
        ])

    add_section_header(doc, "Domain Fit")
    add_bullet_list(doc, [
        "Built Polymarket prediction bot from scratch - direct prediction market domain expertise",
        "Ran paid community with multiple pricing tiers, conversion optimization, and retention analysis",
        "Active prediction market participant familiar with Kalshi's regulatory position and competitive landscape",
        "SQL-native analyst comfortable building dashboards, attribution models, and experimentation frameworks",
    ])

    add_section_header(doc, "Technical Skills")
    add_skills_line(doc, "Analysis & SQL", SKILLS["data_analysis"])
    add_skills_line(doc, "Experimentation", SKILLS["stats"])
    add_skills_line(doc, "Visualization", SKILLS["viz"])
    add_skills_line(doc, "Market Data", SKILLS["market"])

    add_section_header(doc, "Education")
    add_bullet_list(doc, EDUCATION)

    path = os.path.join(OUTPUT_DIR, "Damilare_Kehinde-David_Resume_Kalshi_Growth.docx")
    doc.save(path)
    print(f"  Saved: {path}")


# ── Resume 3: Airwallex - Director, Growth Content ─────────────────────────

def build_airwallex():
    doc = make_doc()
    add_name(doc)
    add_subtitle(doc, "Growth Content Strategy  |  AI Content Operations  |  Multi-Platform Distribution")
    add_contact(doc)

    add_section_header(doc, "Profile")
    add_profile(doc,
        "Content strategist who builds AI-powered content engines at scale. Designed and operates "
        "ContentBrain, an automated system producing 60+ posts/month across 5 platforms with "
        "competitor scraping, trend detection, and performance tracking built in. Combines content "
        "operations expertise with data-driven optimization, B2B SaaS experience, and a deep "
        "understanding of how AI tools transform content production workflows."
    )

    add_section_header(doc, "Content and Growth Experience")
    add_experience_block(doc,
        "ContentBrain - AI Content Engine", "Quivira Brand", "2025 - Present",
        [
            "Built ContentBrain: AI-powered content operations system producing 60+ posts/month across X, LinkedIn, TikTok, Instagram, YouTube",
            "Automated competitor scraping, trend detection, hook analysis, and content performance tracking pipelines",
            "Multi-platform content strategy with platform-specific optimization (format, tone, timing, hooks)",
            "AI tools integration: Claude API for writing, fal.ai for visual generation, automated scheduling and distribution",
            "Content performance analytics: engagement tracking, A/B testing on hooks, conversion attribution from content to subscribers",
        ])

    add_section_header(doc, "Experience")

    add_experience_block(doc,
        EXP_SUBSCRIPTION["title"], EXP_SUBSCRIPTION["org"], EXP_SUBSCRIPTION["date"],
        [
            "Content-driven acquisition funnel: tracked conversion from content impressions to paid subscribers across tiers",
            "Subscriber LTV, cohort retention, and referral attribution informing content strategy decisions",
            "Content performance dashboards in PostgreSQL measuring ROI per platform and content type",
        ])

    add_experience_block(doc,
        EXP_SIGNAL["title"], EXP_SIGNAL["org"], EXP_SIGNAL["date"],
        [
            "Data-driven content: signal performance data feeds content creation (market analysis posts, educational threads)",
            "Full pipeline: ingestion, computation, scoring, filtering - demonstrates content operations at technical depth",
        ])

    add_experience_block(doc,
        EXP_PHARMA["title"], EXP_PHARMA["org"], EXP_PHARMA["date"],
        [
            "B2B SaaS product experience: Super Admin dashboard, subscription lifecycle, multi-tenant analytics",
            "Understands B2B buyer journey and how content drives enterprise pipeline",
        ])

    add_section_header(doc, "Technical Skills")
    add_skills_line(doc, "Content Tools", "Claude API, fal.ai, Typefully, Notion, automated distribution pipelines")
    add_skills_line(doc, "Analytics", SKILLS["data_analysis"])
    add_skills_line(doc, "SEO & Performance", "Content performance analytics, A/B testing, cohort analysis, attribution modeling")
    add_skills_line(doc, "Engineering", SKILLS["eng"])

    add_section_header(doc, "Education")
    add_bullet_list(doc, EDUCATION)

    path = os.path.join(OUTPUT_DIR, "Damilare_Kehinde-David_Resume_Airwallex_Content.docx")
    doc.save(path)
    print(f"  Saved: {path}")


# ── Resume 4: Lido Finance - Institutional Growth Lead USA ──────────────────

def build_lido():
    doc = make_doc()
    add_name(doc)
    add_subtitle(doc, "Institutional DeFi Growth  |  Liquid Staking  |  Quantitative Research")
    add_contact(doc)

    add_section_header(doc, "Profile")
    add_profile(doc,
        "DeFi-native analyst and community builder with deep protocol knowledge across L1/L2 ecosystems, "
        "DEX infrastructure, lending protocols, and liquid staking mechanics. Built trading signal "
        "infrastructure scanning Binance futures and forex markets, and runs a paid trading community "
        "with institutional-grade analytics. Understands how DeFi ecosystem participants interconnect "
        "and the value proposition of liquid staking for institutional allocators."
    )

    add_section_header(doc, "Experience")

    add_experience_block(doc,
        EXP_SIGNAL["title"], EXP_SIGNAL["org"], EXP_SIGNAL["date"],
        [
            "Built institutional-grade signal infrastructure: scans all Binance futures + 8 forex pairs across 3 timeframes",
            "Full DeFi data pipeline: ingestion, indicator computation, Bayesian scoring, threshold filtering",
            "All signals logged to PostgreSQL with outcome tracking, statistical analysis on win rates - demonstrates quantitative rigor",
            "Stack: Python (pandas, NumPy), SQL, REST APIs, Docker, Railway",
        ])

    add_experience_block(doc,
        EXP_POLYMARKET["title"], EXP_POLYMARKET["org"], EXP_POLYMARKET["date"],
        [
            "Designed 7-factor Bayesian scoring system for calibrated probability estimates - quantitative research depth",
            "Factor analysis and weight auditing against historical resolution data",
            "Experience with on-chain event contracts, market microstructure, and probability calibration",
        ])

    add_experience_block(doc,
        EXP_SUBSCRIPTION["title"], EXP_SUBSCRIPTION["org"], EXP_SUBSCRIPTION["date"],
        [
            "Account management: subscriber LTV, churn reduction, cohort retention, tier-based relationship management ($49/wk to $1,499 lifetime)",
            "Referral partnership program (30% commission) - demonstrates stakeholder and partnership management",
            "Signal performance dashboards enabling thought leadership content and institutional-grade reporting",
        ])

    add_experience_block(doc,
        EXP_PHARMA["title"], EXP_PHARMA["org"], EXP_PHARMA["date"],
        [
            "B2B stakeholder management: multi-tenant dashboard serving pharmacy owners, admins, and billing teams",
            "Revenue analytics and subscription lifecycle tracking across diverse stakeholder groups",
        ])

    add_section_header(doc, "Domain Fit")
    add_bullet_list(doc, [
        "Deep DeFi protocol knowledge: L1/L2 ecosystems, DEX infrastructure, lending protocols, bridges, liquid staking mechanics",
        "Built trading signal infrastructure processing real-time market data - understands institutional data requirements",
        "Runs paid trading community with institutional-grade analytics, performance tracking, and stakeholder reporting",
        "Understands liquid staking value proposition for institutional allocators: yield optimization, capital efficiency, protocol composability",
    ])

    add_section_header(doc, "Technical Skills")
    add_skills_line(doc, "DeFi & Markets", SKILLS["market"])
    add_skills_line(doc, "Data Analysis", SKILLS["data_analysis"])
    add_skills_line(doc, "Statistics", SKILLS["stats"])
    add_skills_line(doc, "Engineering", SKILLS["eng"])

    add_section_header(doc, "Education")
    add_bullet_list(doc, EDUCATION)

    path = os.path.join(OUTPUT_DIR, "Damilare_Kehinde-David_Resume_Lido_Institutional.docx")
    doc.save(path)
    print(f"  Saved: {path}")


# ── Main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating tailored resumes...\n")
    build_primeintellect()
    build_kalshi()
    build_airwallex()
    build_lido()
    print(f"\nDone. All 4 resumes saved to:\n  {OUTPUT_DIR}")
